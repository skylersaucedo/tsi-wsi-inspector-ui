"""
Make inspection report from command
"""
import os
import pandas as pd
import cv2
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
import sys

from docx import Document
from docx.shared import Inches
from matplotlib.backends.backend_pdf import PdfPages

def make_pdf_report(defect_log_path, logo_file_path, save_pdf_path):
    """
    use matplotlib to gen a structured pdf page

    """
    # Define the number of images per page
    n_images_per_page = 4

    # Define the size of the figure and the size of each subplot
    figsize = (8.5, 11)
    subplot_size = (5, 3)

    # Define the number of rows and columns of subplots per page
    n_rows_per_page = 2
    n_cols_per_page = 2

    # Define the number of pages

    df = pd.read_csv(defect_log_path)

    print(df)

    n_pages = int(np.ceil(len(df) / n_images_per_page))

    print('number of pages: ', n_pages)


    # Create a new PDF file

    svpath = r'C:\Users\TSI\Desktop\boxscans-jan3\multipage.pdf'

    with PdfPages(svpath) as pdf:
        for i in range(n_pages):
            # Create a new figure
            fig = plt.figure(figsize=figsize)

            # Add subplots to the figure
            for j in range(n_images_per_page):

                z = i * n_images_per_page + j

                if z < len(df):

                    # add img stuff here

                    ax = fig.add_subplot(n_rows_per_page, n_cols_per_page, j + 1)

                    #ax.imshow(images[i * n_images_per_page + j])

                    datetime = df.iloc[z]['datetime']
                    image_path = df.iloc[z]['img_name']
                    notes = df.iloc[z]['notes']
                    x = df.iloc[z]['loc_x']
                    y = df.iloc[z]['loc_y']
                    label = df.iloc[z]['defect']
                    def_h = df.iloc[z]['def_h']
                    def_w = df.iloc[z]['def_w']
                    image_index = df.iloc[z]['img_idx']
                    pipe_ID = df.iloc[z]['pipe_ID']
                    inspector = df.iloc[z]['inspector']

                    # open img, find defect
                    img = cv2.imread(image_path)    
                    defect_r = img[y:y+def_h,x:x+def_w,:]

                    defect = cv2.cvtColor(defect_r, cv2.COLOR_BGR2RGB)


                    ax.imshow(defect)

                    textstr = '\n'.join(('datetime: ', datetime, 
                                         'datetime: ', datetime,
                                         'datetime: ', datetime,
                                         ))
                    
                    ax.text(0,1,textstr, fontsize=14, verticalalignment='top')
                    
                    # ax.text(0.0,0.9, "Image Name: " + image_path)
                    # ax.text(0.0,0.8, "Datetime: " + datetime)
                    # ax.text(0.0,0.7, "Pipe ID: " + pipe_ID)
                    # ax.text(0.0,0.6, "Inspector: " + inspector)
                    # ax.text(0.0,0.5, "Defect: " + label)
                    # ax.text(0.0,0.4, "Notes: " + str(notes))





                    # ax.set_xticks([])
                    # ax.set_yticks([])
                    ax.set_xlabel('Note ' + str(i * n_images_per_page + j + 1))

        


            # Save the figure to the PDF file
            pdf.savefig(bbox_inches='tight')

            # Close the figure
            plt.close(fig)


    # # create a sample figure with two subplots
    # fig, axs = plt.subplots(nrows=len(df)+1, ncols=2, figsize=(20, int(3*len(df))))

    # # add TSI logo as offset image

    # logo_r = cv2.imread(logo_file_path)
    # logo = cv2.cvtColor(logo_r, cv2.COLOR_BGR2RGB)

    # for j in range(2):
        
    #     axs[0,j].imshow(logo)
    #     axs[0,j].tick_params(axis='both', which='both', length=0, labelsize=0)
    #     # remove the border
    #     axs[0,j].spines['top'].set_visible(False)
    #     axs[0,j].spines['right'].set_visible(False)
    #     axs[0,j].spines['bottom'].set_visible(False)
    #     axs[0,j].spines['left'].set_visible(False)

    # for i, row in df.iterrows():
        
    #     datetime = row['datetime']
    #     image_path = row['img_name']
    #     notes = row['notes']
    #     x = row['loc_x']
    #     y = row['loc_y']
    #     label = row['defect']
    #     def_h = row['def_h']
    #     def_w = row['def_w']
    #     image_index = row['img_idx']
    #     pipe_ID = row['pipe_ID']
    #     inspector = row['inspector']

    #     # open img, find defect
    #     img = cv2.imread(image_path)    
    #     defect_r = img[y:y+def_h,x:x+def_w,:]

    #     defect = cv2.cvtColor(defect_r, cv2.COLOR_BGR2RGB)


    #     axs[i+1,0].imshow(defect)
        
    #     axs[i+1,1].text(0.0,0.9, "Image Name: " + image_path)
    #     axs[i+1,1].text(0.0,0.8, "Datetime: " + datetime)
    #     axs[i+1,1].text(0.0,0.7, "Pipe ID: " + pipe_ID)
    #     axs[i+1,1].text(0.0,0.6, "Inspector: " + inspector)
    #     axs[i+1,1].text(0.0,0.5, "Defect: " + label)
    #     axs[i+1,1].text(0.0,0.4, "Notes: " + str(notes))

    #     # remove ticks
    #     axs[i+1,0].tick_params(axis='both', which='both', length=0, labelsize=0)
    #     axs[i+1,1].tick_params(axis='both', which='both', length=0, labelsize=0)
    
    #     # remove the border
    #     axs[i+1,0].spines['top'].set_visible(False)
    #     axs[i+1,0].spines['right'].set_visible(False)
    #     axs[i+1,0].spines['bottom'].set_visible(False)
    #     axs[i+1,0].spines['left'].set_visible(False)
        
    #     # remove the border
    #     axs[i+1,1].spines['top'].set_visible(False)
    #     axs[i+1,1].spines['right'].set_visible(False)
    #     axs[i+1,1].spines['bottom'].set_visible(False)
    #     axs[i+1,1].spines['left'].set_visible(False)



    # # create a PdfPages object to save the plots to a PDF file
    # with PdfPages(save_pdf_path) as pdf:
    #     # save each subplot to the PDF file
    #     pdf.savefig(fig)

    # plt.show()




def make_docx_report(defect_log_path, logo_file_path, save_pdf_path):
    """
    uses python-docx to write a docx file
    """
    df = pd.read_csv(defect_log_path)

    print(df)

    # create a document
    document = Document()

    # add TSI logo as offset image

    logo_r = cv2.imread(logo_file_path)
    logo = cv2.cvtColor(logo_r, cv2.COLOR_BGR2RGB)

    for i, row in df.iterrows():

        if i % 4 == 0:
            document.add_page_break()
        
        datetime = row['datetime']
        image_path = row['img_name']
        notes = row['notes']
        x = row['loc_x']
        y = row['loc_y']
        label = row['defect']
        def_h = row['def_h']
        def_w = row['def_w']
        image_index = row['img_idx']
        pipe_ID = row['pipe_ID']
        inspector = row['inspector']

        # open img, find defect
        img = cv2.imread(image_path)    
        #defect_r = img[y:y+def_h,x:x+def_w,:]

        #defect = cv2.cvtColor(defect_r, cv2.COLOR_BGR2RGB)

        
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(2.5)) # too large



    document.save(r'C:\Users\TSI\Desktop\boxscans-jan3\testreport.docx')




def make_report_from_logs(defect_log_path, logo_file_path, save_pdf_path):

    df = pd.read_csv(defect_log_path)

    print(df)

    # create a sample figure with two subplots
    fig, axs = plt.subplots(nrows=len(df)+1, ncols=2, figsize=(20, int(3*len(df))))

    # add TSI logo as offset image

    logo_r = cv2.imread(logo_file_path)
    logo = cv2.cvtColor(logo_r, cv2.COLOR_BGR2RGB)

    for j in range(2):
        
        axs[0,j].imshow(logo)
        axs[0,j].tick_params(axis='both', which='both', length=0, labelsize=0)
        # remove the border
        axs[0,j].spines['top'].set_visible(False)
        axs[0,j].spines['right'].set_visible(False)
        axs[0,j].spines['bottom'].set_visible(False)
        axs[0,j].spines['left'].set_visible(False)

    for i, row in df.iterrows():
        
        datetime = row['datetime']
        image_path = row['img_name']
        notes = row['notes']
        x = row['loc_x']
        y = row['loc_y']
        label = row['defect']
        def_h = row['def_h']
        def_w = row['def_w']
        image_index = row['img_idx']
        pipe_ID = row['pipe_ID']
        inspector = row['inspector']

        # open img, find defect
        img = cv2.imread(image_path)    
        defect_r = img[y:y+def_h,x:x+def_w,:]

        defect = cv2.cvtColor(defect_r, cv2.COLOR_BGR2RGB)


        axs[i+1,0].imshow(defect)
        
        axs[i+1,1].text(0.0,0.9, "Image Name: " + image_path)
        axs[i+1,1].text(0.0,0.8, "Datetime: " + datetime)
        axs[i+1,1].text(0.0,0.7, "Pipe ID: " + pipe_ID)
        axs[i+1,1].text(0.0,0.6, "Inspector: " + inspector)
        axs[i+1,1].text(0.0,0.5, "Defect: " + label)
        axs[i+1,1].text(0.0,0.4, "Notes: " + str(notes))

        # remove ticks
        axs[i+1,0].tick_params(axis='both', which='both', length=0, labelsize=0)
        axs[i+1,1].tick_params(axis='both', which='both', length=0, labelsize=0)
    
        # remove the border
        axs[i+1,0].spines['top'].set_visible(False)
        axs[i+1,0].spines['right'].set_visible(False)
        axs[i+1,0].spines['bottom'].set_visible(False)
        axs[i+1,0].spines['left'].set_visible(False)
        
        # remove the border
        axs[i+1,1].spines['top'].set_visible(False)
        axs[i+1,1].spines['right'].set_visible(False)
        axs[i+1,1].spines['bottom'].set_visible(False)
        axs[i+1,1].spines['left'].set_visible(False)



    # create a PdfPages object to save the plots to a PDF file
    with PdfPages(save_pdf_path) as pdf:
        # save each subplot to the PDF file
        pdf.savefig(fig)

    plt.show()



if __name__ == "__main__":
    
    root = os.getcwd()
    print('root: ', root)

    project_path = sys.argv[1]

    print('project path: ', project_path)

    #save_pdf_path = r'C:\Users\TSI\source\repos\tsi-wsi-inspector-ui\bbox_maker\inspection_report.pdf'
    save_pdf_path = project_path + "\\" + "Metadata" + "\\" + "inspection_report.pdf"

    defect_log_path = r'C:\Users\TSI\source\repos\tsi-wsi-inspector-ui\bbox_maker\defectslog.txt'
    logo_file_path = r'C:\Users\TSI\source\repos\tsi-wsi-inspector-ui\bbox_maker\logoTUBES.png'

    make_report_from_logs(defect_log_path, logo_file_path, save_pdf_path)
    
    #make_docx_report(defect_log_path, logo_file_path, save_pdf_path)

    make_pdf_report(defect_log_path, logo_file_path, save_pdf_path)
